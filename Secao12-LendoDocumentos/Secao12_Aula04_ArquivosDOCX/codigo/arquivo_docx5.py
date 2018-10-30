import docx

doc = docx.Document('ArquivoWord3.docx')

paragrafos = len(doc.paragraphs)
contador = 0

while contador < paragrafos:
    print(doc.paragraphs[contador].style.name)
    if doc.paragraphs[contador].style.name == 'Heading 4':
        doc.paragraphs[contador].style.name = 'Normal'
    contador += 1

doc.save('arquivo_salvo.docx')