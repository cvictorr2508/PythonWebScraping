import docx

doc = docx.Document('ArquivoWord2.docx')

total_runs = len(doc.paragraphs[0].runs)
contador = 0

while contador < total_runs:
    print(doc.paragraphs[0].runs[contador].text)
    contador += 1